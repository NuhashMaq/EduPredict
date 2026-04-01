from __future__ import annotations

from pathlib import Path

from docx import Document


def add_code_block(doc: Document, lines: list[str]) -> None:
    # Word has no native fenced-code; we approximate with a monospace paragraph per line.
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    in_code = False
    code_lines: list[str] = []

    for raw in text.splitlines():
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, code_lines)
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("1. ") or line[:3].isdigit() and line[3:5] == ". ":
            # Very small heuristic for ordered lists
            doc.add_paragraph(line.split(". ", 1)[1].strip(), style="List Number")
        else:
            doc.add_paragraph(line)

    if in_code and code_lines:
        add_code_block(doc, code_lines)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main() -> None:
    root = Path(__file__).resolve().parent
    md = root / "CSE_3200_Project_Report_EduPredict.md"
    out = root / "CSE_3200_Project_Report_EduPredict.docx"

    if not md.exists():
        raise SystemExit(f"Missing input markdown: {md}")

    md_to_docx(md, out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
